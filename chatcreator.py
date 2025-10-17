import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import docx

# Percorsi dei file
input_excel_file = 'chat_with_transcriptions.xlsx'  # Questo file deve essere generato dal programma di Cellebrite
output_docx_file = 'chat_simulation.docx'
attachment_base_path = 'instant_messages/WhatsApp/1'

# Leggi il file Excel
df = pd.read_excel(input_excel_file)

# Crea un nuovo documento Word
doc = Document()

# Funzione per creare un hyperlink in un paragrafo
def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a w:r element and a new w:rPr element
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Add color if it is given
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'none')
        rPr.append(u)

    new_run.append(rPr)
    # Create a new w:t element and add the text
    new_run.append(OxmlElement('w:t'))
    new_run[-1].text = text

    # Add the w:r element to the w:hyperlink element
    hyperlink.append(new_run)

    # Add the hyperlink to the paragraph
    paragraph._element.append(hyperlink)

    return paragraph

# Funzione per creare una bolla di chat
def create_chat_bubble(doc, message, timestamp, incoming=True, is_attachment=False, attachment_url=None):
    bubble = doc.add_paragraph()
    if is_attachment and attachment_url:
        bubble = add_hyperlink(bubble, attachment_url, message, '0000FF', True)
    else:
        run = bubble.add_run(message)

    bubble.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT if incoming else WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Aggiungi il timestamp
    timestamp_run = bubble.add_run(f"\n{timestamp}")
    timestamp_run.font.size = Pt(8)
    timestamp_run.font.color.rgb = RGBColor(128, 128, 128)

    # Imposta stile
    for r in bubble.runs:
        r.font.size = Pt(12)

    # Crea una bolla
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), "E0E0E0" if incoming else "ADD8E6")
    bubble._element.get_or_add_pPr().append(shading_elm)

    # Imposta i margini della bolla
    if incoming:
        bubble.paragraph_format.left_indent = Inches(0.25)
        bubble.paragraph_format.right_indent = Inches(2)
    else:
        bubble.paragraph_format.left_indent = Inches(2)
        bubble.paragraph_format.right_indent = Inches(0.25)

# Aggiungi separatore per distinguere tra i messaggi
def add_separator(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run('-' * 50)
    run.font.size = Pt(5)
    run.font.bold = True
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Itera attraverso le righe del DataFrame e aggiungi i messaggi al documento Word
for index, row in df.iterrows():
    direction = row['Direction']
    # Controllo se il valore è una stringa, altrimenti lo converto in stringa
    if isinstance(direction, str):
        incoming = direction.lower() == 'incoming'
    else:
        # Gestione del caso in cui il valore non sia una stringa o non ci sia
        incoming = False  # imposto a False per evitare errori

    message = row['Body']
    attachment = row['attachments']
    transcription = row['Transcription']
    timestamp = row['Timestamp-Time']

    if pd.notna(message) and message.strip() != '':
        create_chat_bubble(doc, message, timestamp, incoming)

    if pd.notna(attachment) and str(attachment).strip() != '':
        # Estrarre il nome del file e costruire il percorso relativo
        attachment_file = str(attachment).replace("Attachment: ", "").strip()
        attachment_url = os.path.join(attachment_base_path, attachment_file)
        create_chat_bubble(doc, f"Attachment: {attachment_file}", timestamp, incoming, is_attachment=True, attachment_url=attachment_url)

    if pd.notna(transcription) and transcription.strip() != '':
        create_chat_bubble(doc, f"Transcription: {transcription}", timestamp, incoming)

    # Aggiungi un separatore dopo ogni messaggio
    add_separator(doc)

# Salva il documento Word
doc.save(output_docx_file)

print(f"Chat simulation saved in {output_docx_file}")
