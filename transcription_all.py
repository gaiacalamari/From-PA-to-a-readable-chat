import os
import azure.cognitiveservices.speech as speechsdk
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docxcompose.composer import Composer
from tqdm import tqdm

# Recupera le credenziali dalle variabili d'ambiente
subscription_key = os.getenv('MARCO_AZURE_API_KEY')
region = 'westeurope'

if not subscription_key:
    raise ValueError("La variabile d'ambiente MARCO_AZURE_API_KEY non è impostata.")

lingua = 'it-IT'

# Funzione per trascrivere un singolo file audio
def transcribe_file(audio_path):
    global continuous_recognition_done
    transcription_results = []

    def recognized_handler(evt):
        transcription_results.append(evt.result.text)

    def session_stopped_handler(evt):
        global continuous_recognition_done
        continuous_recognition_done = True

    def canceled_handler(evt):
        if evt.result.reason == speechsdk.ResultReason.Canceled:
            cancellation_details = evt.result.cancellation_details
            print(f"Errore: {cancellation_details.reason}, Dettagli: {cancellation_details.error_details}")
        global continuous_recognition_done
        continuous_recognition_done = True

    try:
        audio_config = speechsdk.audio.AudioConfig(filename=audio_path)
        speech_recognizer = speechsdk.SpeechRecognizer(speech_config=speech_config, audio_config=audio_config)
        speech_config.set_profanity(speechsdk.ProfanityOption.Raw)  # Disabilita la censura delle parolacce

        # Collega i gestori degli eventi
        speech_recognizer.recognized.connect(recognized_handler)
        speech_recognizer.session_stopped.connect(session_stopped_handler)
        speech_recognizer.canceled.connect(canceled_handler)

        continuous_recognition_done = False

        speech_recognizer.start_continuous_recognition()

        # Attendi che la trascrizione sia completata
        while not continuous_recognition_done:
            pass

        speech_recognizer.stop_continuous_recognition()

        return ' '.join(transcription_results)

    except Exception as e:
        print(f"Errore durante la trascrizione: {e}")
        return None

# Funzione per unire i documenti Word
def merge_documents(documents, output_path):
    master = Document()
    composer = Composer(master)

    for doc in documents:
        composer.append(Document(doc))

    composer.save(output_path)
    print(f"Tutti i documenti sono stati uniti in: {output_path}")

# Funzione per leggere il file di log e ottenere i file già trascritti
def read_log(log_file):
    if os.path.exists(log_file):
        with open(log_file, 'r') as file:
            return set(file.read().splitlines())
    return set()

# Funzione per aggiornare il file di log
def update_log(log_file, filename):
    with open(log_file, 'a') as file:
        file.write(filename + '\n')

# Funzione per trascrivere i file audio in una directory e salvare i risultati in documenti Word
def transcribe_directory_to_docx(directory, output_file_prefix, batch_size=100, log_file='transcription_log.txt'):
    doc = Document()
    file_count = 0
    batch_count = 0
    batch_files = []
    total_files = len([name for name in os.listdir(directory) if name.endswith('.wav')])

    completed_files = read_log(log_file)

    for filename in tqdm(os.listdir(directory), total=total_files, desc="Trascrizione dei file audio"):
        if filename.endswith('.wav') and filename not in completed_files:
            audio_path = os.path.join(directory, filename)
            print(f"Trascrizione del file: {audio_path}")
            transcription = transcribe_file(audio_path)
            if transcription:
                doc.add_paragraph(f"File: {filename}")
                doc.add_paragraph(f"Testo trascritto: {transcription}")
                doc.add_paragraph("-" * 50)  # Separatore
                print(f"Trascrizione completata per il file: {filename}")
                update_log(log_file, filename)
            else:
                print(f"Trascrizione fallita per il file: {filename}")

            file_count += 1

            # Salva il documento ogni batch_size trascrizioni
            if file_count % batch_size == 0:
                batch_filename = f"{output_file_prefix}_batch_{batch_count}.docx"
                doc.save(batch_filename)
                batch_files.append(batch_filename)
                doc = Document()  # Crea un nuovo documento per il prossimo batch
                batch_count += 1
                print(f"Batch {batch_count} salvato come {batch_filename}")
                os.startfile(batch_filename)  # Apri il documento Word

    # Salva l'ultimo batch se ci sono trascrizioni rimanenti
    if file_count % batch_size != 0:
        batch_filename = f"{output_file_prefix}_batch_{batch_count}.docx"
        doc.save(batch_filename)
        batch_files.append(batch_filename)
        print(f"Ultimo batch salvato come {batch_filename}")
        os.startfile(batch_filename)  # Apri il documento Word

    # Unisci tutti i batch in un unico documento finale
    merge_documents(batch_files, f"{output_file_prefix}_final.docx")

# Configurazione del servizio di riconoscimento vocale
speech_config = speechsdk.SpeechConfig(subscription=subscription_key, region=region)
speech_config.speech_recognition_language = lingua

# Percorsi delle directory
audio_directory = 'D:/P/_Audio trascrizione/My_audio'
output_file_prefix = 'D:/P/_Audio trascrizione/Trascrizione/trascrizione.docx'
log_file = 'D:/P/_Audio trascrizione/Trascrizione/transcription_log.txt'

# Esegui la trascrizione su tutti i file nella directory e salva in documenti Word
transcribe_directory_to_docx(audio_directory, output_file_prefix, log_file=log_file)
